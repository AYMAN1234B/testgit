import docx
import io
from typing import Union

class DocxProcessor:
    def __init__(self):
        pass
    
    def extract_text(self, docx_content: Union[bytes, io.BytesIO]) -> str:
        """Extract text from Word document (.docx)"""
        try:
            if isinstance(docx_content, bytes):
                docx_content = io.BytesIO(docx_content)
            
            # Reset stream position
            docx_content.seek(0)
            
            # Open the document
            doc = docx.Document(docx_content)
            
            # Extract text from all paragraphs
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            # Join all text with proper spacing
            full_text = "\n".join(text_content)
            
            # Clean and return the text
            return self._clean_text(full_text)
            
        except Exception as e:
            raise Exception(f"Could not extract text from Word document: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        import re
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Split into lines for processing
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            
            # Skip very short lines that are likely artifacts
            if len(line) < 2:
                continue
                
            # Skip lines with only special characters
            if re.match(r'^[^\w\s]+$', line):
                continue
            
            cleaned_lines.append(line)
        
        # Rejoin text
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Final cleanup
        cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
        cleaned_text = cleaned_text.strip()
        
        return cleaned_text
    
    def validate_docx(self, docx_content: Union[bytes, io.BytesIO]) -> dict:
        """Validate Word document and return metadata"""
        try:
            if isinstance(docx_content, bytes):
                docx_content = io.BytesIO(docx_content)
            
            # Reset stream position
            docx_content.seek(0)
            
            # Try to open the document
            doc = docx.Document(docx_content)
            
            # Count paragraphs and tables
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
            table_count = len(doc.tables)
            
            # Get document properties
            core_properties = doc.core_properties
            
            metadata = {
                "valid": True,
                "paragraph_count": paragraph_count,
                "table_count": table_count,
                "title": core_properties.title or "",
                "author": core_properties.author or "",
                "subject": core_properties.subject or "",
                "created": str(core_properties.created) if core_properties.created else "",
                "modified": str(core_properties.modified) if core_properties.modified else "",
            }
            
            return metadata
            
        except Exception as e:
            return {
                "valid": False,
                "error": str(e),
                "paragraph_count": 0,
                "table_count": 0
            }